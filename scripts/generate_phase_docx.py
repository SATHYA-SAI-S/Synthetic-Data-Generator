"""Generate a Phase-by-Phase Explanatory DOCX for the DP Healthcare Framework."""
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = os.path.join("docs", "Phase_Explanatory_Document.docx")
doc = Document()
st = doc.styles["Normal"]
st.font.name = "Calibri"
st.font.size = Pt(11)
BLUE = RGBColor(0x1F, 0x3A, 0x5F)

def h(text, level=1):
    x = doc.add_heading(text, level=level)
    for r in x.runs:
        r.font.color.rgb = BLUE

def p(text, bold=False):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    return para

def b(text):
    doc.add_paragraph(text, style="List Bullet")

def tbl(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, hd in enumerate(headers):
        t.rows[0].cells[i].text = hd
        for r in t.rows[0].cells[i].paragraphs[0].runs:
            r.bold = True
    for ri, row in enumerate(rows, 1):
        for ci, val in enumerate(row):
            t.rows[ri].cells[ci].text = str(val)
    return t

# Title Page
tl = doc.add_heading("Privacy-Preserving Synthetic Healthcare Data\nGeneration Framework", 0)
tl.alignment = WD_ALIGN_PARAGRAPH.CENTER
for r in tl.runs:
    r.font.color.rgb = BLUE
sp = doc.add_paragraph()
sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sp.add_run("Phase-by-Phase Explanatory Document\nPhases 1-10\nDate: 2026-08-19 | Version 1.0")
run.bold = True
run.font.size = Pt(14)
doc.add_page_break()

# Table of Contents
h("Table of Contents", 1)
toc = [
    "1.  Overview & Architecture",
    "2.  Phase 1: Project Setup & Environment",
    "3.  Phase 2: Environment Verification",
    "4.  Phase 3: Data Profiling & Preprocessing",
    "5.  Phase 4: Diffusion Model Baseline",
    "6.  Phase 5: Differential Privacy (DP-SGD)",
    "7.  Phase 6: Hyperparameter Sweep",
    "8.  Phase 7: DP-Diffusion Training & Generation",
    "9.  Phase 8: Evaluation (Utility & Privacy)",
    "10. Phase 9: Red-Team Validation",
    "11. Phase 10: Final Report & Deliverables",
    "12. Appendix: File Map",
]
for item in toc:
    p(item)
doc.add_page_break()

# ── 1. OVERVIEW ──
h("1. Overview & Architecture", 1)
p("This framework generates synthetic healthcare data that preserves the statistical properties of real patient records while providing a formal differential privacy (DP) guarantee. It combines a diffusion model for high-fidelity generation with DP-SGD for privacy protection, and a preprocessing pipeline that handles HIPAA identifiers, missingness, and mixed data types.")
h("1.1 Four-Layer Architecture", 2)
b("Layer 1 - Config (src/config/schema.py): Frozen Pydantic models; source of truth for all thresholds.")
b("Layer 2 - Profiling (src/profiling/): DatasetProfiler; HIPAA detection; dtype inference; missingness analysis.")
b("Layer 3 - Preprocessing (src/preprocessing/): MissingnessHandler; invertible encoders and scalers; PreprocessingPipeline orchestrator.")
b("Layer 4 - Registry (src/registry/): FileSchemaRegistry; atomic versioned persistence with SHA-256 integrity.")
h("1.2 Key Design Principles", 2)
b("Dependency Injection: all components injected via constructors, enabling testability and swap-ability.")
b("Invertible Transforms: every encoder/scaler implements inverse_transform for exact round-trip decoding.")
b("I/O at the Boundary: file operations only at the outermost edge (fit_transform_from_file).")
b("Frozen Config: immutable Pydantic models prevent accidental threshold mutation.")
b("Atomic Writes: registry and budget state use temp-file + rename to survive session kills.")
doc.add_page_break()

# ── 2. PHASE 1 ──
h("2. Phase 1: Project Setup & Environment", 1)
p("Objective: Establish project structure, define configuration, and set up the development environment with all required dependencies.")
h("2.1 Key Configuration Components", 2)
tbl(["Component", "Purpose"], [
    ["PipelineConfig", "Root config object composed of all sub-configs."],
    ["CardinalityConfig", "low_card_max=15, near_identifier_ratio=0.95, rare_category_min_freq=10."],
    ["MissingnessConfig", "drop_if_missing_above=0.80, inject_indicator_above=0.01, structural_correlation_threshold=0.40."],
    ["SmallNConfig", "small_n_threshold=500, minimum_viable_n=100."],
    ["DiffusionConfig", "num_timesteps=1000, beta_start=1e-4, beta_end=0.02, hidden_dims=[256,256,256]."],
    ["TrainingConfig", "batch_size=256, lr=2e-4, epochs=50, gpu_budget_hours=30."],
    ["PrivacyConfig", "target_epsilon=1.0, target_delta=1e-5, max_grad_norm=1.0."],
])
h("2.2 Main Dependencies", 2)
tbl(["Package", "Version", "Purpose"], [
    ["torch", "2.3.1 (cu118)", "Deep-learning framework for the diffusion model."],
    ["opacus", "1.5.2", "DP-SGD, per-sample gradients, and RDP accountant."],
    ["pandas", "2.2.2", "Data manipulation."],
    ["numpy", "1.26.4", "Numerical computing."],
    ["scikit-learn", "1.5.1", "ML utilities (NearestNeighbors for MIA)."],
    ["scipy", "1.13.1", "Statistical tests (KS, point-biserial, Cramer's V)."],
    ["pydantic", "2.7.4", "Typed configuration schemas (V2 API)."],
    ["joblib", "1.4.2", "Serializer for scalers, encoders, and pipeline state."],
    ["pytest", "8.2.2", "Testing framework."],
])
doc.add_page_break()

# ── 3. PHASE 2 ──
h("3. Phase 2: Environment Verification", 1)
p("Objective: Verify Python version, GPU availability, and Opacus functionality before running DP-SGD training on Kaggle.")
h("3.1 setup_check.py Checks", 2)
b("Python version >= 3.10.")
b("All required packages importable (torch, opacus, pandas, numpy, sklearn, pydantic, scipy, joblib, pyarrow).")
b("CUDA GPU availability and VRAM detection.")
b("Pydantic V2 API functionality.")
b("Opacus DP-SGD smoke test - runs one DP training step and verifies epsilon > 0.")
p("Command:", bold=True)
p("python environment/setup_check.py --report-file setup_report.json")
doc.add_page_break()

# ── 4. PHASE 3 ──
h("4. Phase 3: Data Profiling & Preprocessing", 1)
p("Objective: Profile the raw healthcare dataset, detect HIPAA identifiers, infer dtypes, handle missingness, and transform features into a normalized numeric tensor while preserving full invertibility.")
h("4.1 DatasetProfiler", 2)
b("HIPAA Safe Harbor detection: 18 identifier categories matched with boundary-anchored regex.")
b("Dtype inference: continuous, ordinal, categorical_low, categorical_high, binary, near_identifier, unknown.")
b("Missingness classification: none, MCAR-like, structural, high.")
b("Structural missingness: point-biserial correlation (numeric) or Cramer's V (categorical).")
h("4.2 PreprocessingPipeline", 2)
b("Column dropping: high-missingness, near-identifier, HIPAA, and constant columns.")
b("MissingnessHandler: injects <col>__missing_flag indicator columns; imputes NaN with median/mode.")
b("OneHotEncoder: low-cardinality categoricals with __null__ and __other__ reserved tokens.")
b("FrequencyEncoder: high-cardinality categoricals as frequency-ordered integer indices (embedding-friendly).")
b("StandardScaler / MinMaxScaler / RobustScaler: continuous-variable normalization.")
b("Schema Registry: saves fitted pipeline state for later inverse_transform.")
h("4.3 Round-Trip Contract", 2)
p("Critical invariant: inverse_transform(transform(x)) == x. Every encoder and scaler implements inverse_transform exactly, so generated synthetic data can be decoded back to the original feature space.")
doc.add_page_break()

# ── 5. PHASE 4 ──
h("5. Phase 4: Diffusion Model Baseline", 1)
p("Objective: Implement a baseline (non-DP) diffusion model to establish the generative capability of the framework before DP is added.")
h("5.1 Components", 2)
tbl(["Component", "File", "Description"], [
    ["LinearNoiseSchedule", "src/diffusion/schedule.py", "DDPM linear beta schedule from 1e-4 to 0.02 over 1000 steps."],
    ["MLPDenoiser", "src/diffusion/denoiser.py", "Time-conditioned MLP (Linear->SiLU->LayerNorm) with embedding-based time conditioning."],
    ["forward_diffuse", "src/diffusion/forward_process.py", "Closed-form forward q(x_t|x_0)=sqrt(alpha)*x_0+sqrt(1-alpha)*noise."],
    ["generate_samples", "src/diffusion/sampler.py", "Reverse DDPM sampling with chunked generation to avoid VRAM exhaustion."],
    ["DiffusionTrainer", "src/diffusion/trainer.py", "Non-DP MSE training loop."],
])
h("5.2 Forward Process", 2)
p("x_t = sqrt(alpha_bar_t) * x_0 + sqrt(1 - alpha_bar_t) * epsilon")
h("5.3 Reverse Process (Sampling)", 2)
p("x_{t-1} = (1/sqrt(a_t)) * (x_t - ((1-a_t)/sqrt(1-a_bar_t)) * eps_pred) + sigma_t * z")
doc.add_page_break()

# ── 6. PHASE 5 ──
h("6. Phase 5: Differential Privacy (DP-SGD)", 1)
p("Objective: Add a formal differential-privacy guarantee to the diffusion training using DP-SGD with Opacus.")
h("6.1 DP-SGD Mechanism", 2)
b("Per-sample gradient clipping to a maximum L2 norm C.")
b("Gaussian noise injection: noise ~ N(0, sigma^2 * C^2).")
b("Privacy accounting via an RDP accountant tracking cumulative privacy loss.")
h("6.2 Key Components", 2)
tbl(["Component", "File", "Description"], [
    ["DPTrainer", "src/privacy/dp_trainer.py", "DP-SGD training loop wrapping the denoiser in a GradSampleModule."],
    ["CentralPrivacyAccountant", "src/privacy/accountant.py", "Single RDP accountant; all noise steps are recorded here."],
    ["AdaptiveNoiseSchedule", "src/privacy/adaptive_schedule.py", "Per-timestep sigma: more noise at t=0, less at t=T."],
    ["clip_and_noise_tier", "src/privacy/clip_and_noise.py", "Per-sample clipping + Gaussian noise + accountant record."],
    ["HeuristicRiskTierAssigner", "src/privacy/risk_tier_assigner.py", "Assigns features to Tier1/2/3 based on uniqueness and HIPAA."],
])
h("6.3 Adaptive Noise Schedule", 2)
p("sigma(t) = base_sigma * (0.5 + ratio),  ratio = 1 - t/(T-1)")
h("6.4 Risk Tier Assignment", 2)
b("Tier 1 (Strict): HIPAA-identifier match OR uniqueness > 0.8.")
b("Tier 2 (Moderate): uniqueness > 0.15.")
b("Tier 3 (Loose): low cardinality/uniqueness.")
p("A correlation guard promotes correlated features to the stricter tier to prevent leakage across tiers.")
h("6.5 Critical Fix Found During Audit", 2)
p("The 4th audit discovered that tier_params were extracted from the ORIGINAL denoiser before GradSampleModule wrapping. Because Opacus copies the module, the original parameters never received .grad_sample, causing clip_and_noise_tier to silently no-op - NO DP noise and NO privacy accounting were ever applied. The fix remaps tier_params to the WRAPPED module parameters, restoring the actual DP guarantee.")
doc.add_page_break()

# ── 7. PHASE 6 ──
h("7. Phase 6: Hyperparameter Sweep", 1)
p("Objective: Run a systematic sweep over privacy budgets (epsilon) to explore the privacy-utility tradeoff.")
h("7.1 Sweep Configuration", 2)
tbl(["Parameter", "Values", "Purpose"], [
    ["Epsilon", "[0.1, 1.0, 10.0]", "Privacy budget - lower is more private."],
    ["Base sigma", "15.0 / epsilon", "Higher eps -> lower noise -> better utility."],
    ["Epochs", "50", "Training iterations per epsilon."],
    ["Batch size", "256", "Samples per step."],
    ["Timesteps", "1000", "Diffusion steps."],
    ["Hidden dims", "[256, 256, 256]", "MLP architecture."],
])
h("7.2 ComputeBudgetGuard", 2)
p("Enforces a strict weekly compute budget (default 30 hours) on Kaggle and persists elapsed time to disk so it survives kernel restarts.")
h("7.3 Sweep Outputs", 2)
b("synthetic_eps_{epsilon}.csv - generated synthetic data per epsilon.")
b("sweep_report.json - target epsilon, actual epsilon spent, and loss.")
b("registry_eps_{epsilon}/ - fitted pipeline state.")
b("checkpoints/model_eps_{epsilon}.pt - trained checkpoints.")
doc.add_page_break()

# ── 8. PHASE 7 ──
h("8. Phase 7: DP-Diffusion Training & Generation", 1)
p("Objective: Train the DP-protected diffusion model and generate synthetic healthcare data with formal privacy guarantees.")
h("8.1 End-to-End Pipeline (scripts/reproduce_end_to_end.py)", 2)
b("1. Load the raw UCI diabetes dataset.")
b("2. Profile and preprocess (HIPAA, missingness, encoding, scaling).")
b("3. Assign per-feature privacy risk tiers.")
b("4. Initialize the DP-diffusion components.")
b("5. Train with DP-SGD for N epochs.")
b("6. Generate synthetic samples via reverse diffusion.")
b("7. Inverse-transform to the original feature space.")
b("8. Save synthetic CSV and sweep report.")
h("8.2 Generation", 2)
p("Reverse diffusion runs from pure Gaussian noise through the trained denoiser for the full timesteps, producing samples that follow the learned data distribution.")
doc.add_page_break()

# ── 9. PHASE 8 ──
h("9. Phase 8: Evaluation (Utility & Privacy)", 1)
p("Objective: Evaluate synthetic data quality on two axes - statistical utility and resistance to membership inference.")
h("9.1 UtilityEvaluator", 2)
b("Univariate: KS test for continuous columns; Total Variation Distance for categorical columns.")
b("Bivariate: RMSE between Pearson correlation matrices of real and synthetic data.")
h("9.2 PrivacyEvaluator (D-MIA)", 2)
b("Compute distances from train records to their nearest synthetic neighbor.")
b("Compute distances from holdout records to their nearest synthetic neighbor.")
b("mia_risk_score = max(0, (holdout_dist - train_dist) / holdout_dist), clamped to [0, 1].")
h("9.3 Key Files", 2)
b("src/evaluation/utility_metrics.py - UtilityEvaluator.")
b("src/evaluation/privacy_metrics.py - PrivacyEvaluator.")
b("scripts/run_phase8_evaluation.py - Phase 8 runner.")
doc.add_page_break()

# ── 10. PHASE 9 ──
h("10. Phase 9: Red-Team Validation", 1)
p("Objective: Adversarial validation ensuring the synthetic output does not leak sensitive information or reproduce exact patient rows.")
h("10.1 Validation Checks", 2)
tbl(["Check", "Pass Criteria"], [
    ["HIPAA Leak Detection", "No HIPAA identifier columns present in synthetic data."],
    ["Exact Row Memorization", "0 exact row matches between real and synthetic data."],
    ["Membership Inference", "Low D-MIA risk score."],
])
p("Results are saved to docs/evaluation_results.json with a PASS/FAIL status.")
doc.add_page_break()

# ── 11. PHASE 10 ──
h("11. Phase 10: Final Report & Deliverables", 1)
p("Objective: Consolidate all findings, results, and deliverables into a comprehensive final report and ensure end-to-end reproducibility.")
h("11.1 Deliverables", 2)
tbl(["Deliverable", "Location"], [
    ["Final Capstone Report", "docs/Phase10_Capstone_Report.docx"],
    ["Project Abstracts", "docs/Project_Abstract*.docx"],
    ["Architecture Diagrams", "docs/*.svg"],
    ["Audit Reports", "docs/*_audit_report.md (4 reports)"],
    ["Phase Reports", "docs/phase*.md"],
    ["Evaluation Results", "docs/evaluation_results.json"],
    ["Privacy-Utility Tradeoff", "docs/privacy_utility_tradeoff.png"],
])
h("11.2 Reproducibility", 2)
b("kaggle_runner/run_pipeline.py - autonomous Kaggle runner (clone, install, verify CUDA, run).")
b("scripts/reproduce_end_to_end.py - end-to-end pipeline entry point.")
b("136 automated tests covering all components.")
doc.add_page_break()

# ── 12. APPENDIX ──
h("12. Appendix: File Map", 1)
h("Source Code (src/)", 2)
tbl(["Module", "Files", "Purpose"], [
    ["config/", "schema.py", "Frozen config models."],
    ["profiling/", "base.py, dataset_profiler.py", "Profiling, HIPAA detection, dtype inference."],
    ["preprocessing/", "base.py, encoders.py, missingness.py, pipeline.py, scalers.py", "Preprocessing with invertible transforms."],
    ["privacy/", "accountant.py, adaptive_schedule.py, base.py, clip_and_noise.py, dp_trainer.py, risk_tier_assigner.py", "DP-SGD, accounting, risk tiering."],
    ["diffusion/", "base.py, denoiser.py, forward_process.py, sampler.py, schedule.py, trainer.py", "Diffusion model components."],
    ["evaluation/", "privacy_metrics.py, utility_metrics.py", "Utility & privacy evaluation."],
    ["orchestration/", "gpu_budget_guard.py", "Compute budget enforcement."],
    ["registry/", "base.py, schema_registry.py", "Versioned persistence with SHA-256."],
    ["tests/", "conftest.py, test_*., diffusion/, privacy/, evaluation/", "136 automated tests."],
])
h("Scripts & Runners", 2)
tbl(["File", "Purpose"], [
    ["scripts/reproduce_end_to_end.py", "End-to-end pipeline entry point."],
    ["scripts/reproduce_end_to_end.sh", "Bash entry point."],
    ["scripts/mini_run.py", "Fast mini-run for testing (2k rows)."],
    ["scripts/run_phase8_evaluation.py", "Phase 8 evaluation runner."],
    ["scripts/generate_phase_docx.py", "Generates this document."],
    ["kaggle_runner/run_pipeline.py", "Autonomous Kaggle runner."],
    ["environment/setup_check.py", "Environment verification."],
])

os.makedirs("docs", exist_ok=True)
doc.save(OUT)
print(f"Saved: {OUT}")