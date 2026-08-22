import sys
import re
import glob
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def format_docx(filepath):
    doc = Document(filepath)
    full_text = '\n'.join([p.text for p in doc.paragraphs])
    
    new_doc = Document()
    
    # Add title style formatting
    style = new_doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    for line in full_text.split('\n'):
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('# '):
            h = new_doc.add_heading(line[2:].strip(), level=1)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('## '):
            new_doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith('### '):
            new_doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith('- ') or line.startswith('* '):
            p = new_doc.add_paragraph(style='List Bullet')
            _add_formatted_text(p, line[2:].strip())
        elif line[0].isdigit() and line[1:3] == '. ':
            p = new_doc.add_paragraph(style='List Number')
            _add_formatted_text(p, line)
        else:
            p = new_doc.add_paragraph()
            _add_formatted_text(p, line)
            
    new_doc.save(filepath)

def _add_formatted_text(paragraph, text):
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)

for f in glob.glob('docs/*.docx'):
    if 'Phase1_' in f or 'Phase2_' in f or 'Phase3_' in f or 'Final_' in f or 'Phase10_' in f:
        print(f"Formatting {f}...")
        try:
            format_docx(f)
        except Exception as e:
            print(f"Error formatting {f}: {e}")
